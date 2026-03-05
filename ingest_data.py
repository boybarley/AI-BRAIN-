"""
ingest_data.py — Document ingestion pipeline
Supports: .txt, .md, .pdf, .docx, .csv, .xlsx
Usage:
    python ingest_data.py            # Replace mode (default)
    python ingest_data.py APPEND     # Append to existing index
"""

import os
import sys
import logging
from typing import List

import yaml
from langchain_core.documents import Document

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def load_config() -> dict:
    with open(os.path.join(BASE_DIR, "config.yaml"), "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_documents(data_dir: str) -> List[Document]:
    documents = []

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        return documents

    for filename in sorted(os.listdir(data_dir)):
        filepath = os.path.join(data_dir, filename)
        if not os.path.isfile(filepath):
            continue

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        # Skip hidden / system files
        if filename.startswith(".") or ext == "gitkeep":
            continue

        try:
            # ── TXT / Markdown ──────────────────────────────
            if ext in ("txt", "md"):
                content = None
                for enc in ("utf-8", "latin-1", "cp1252"):
                    try:
                        with open(filepath, "r", encoding=enc) as f:
                            content = f.read()
                        logger.info(f"  TXT: {filename} — loaded ({enc})")
                        break
                    except UnicodeDecodeError:
                        continue

                if content is None:
                    logger.warning(f"  SKIP: {filename} — encoding error")
                    continue

                documents.append(
                    Document(page_content=content, metadata={"source": filename})
                )

            # ── PDF ─────────────────────────────────────────
            elif ext == "pdf":
                from pypdf import PdfReader

                reader = PdfReader(filepath)
                pages = [p.extract_text() or "" for p in reader.pages]
                content = "\n\n".join(pages).strip()

                if not content:
                    logger.warning(f"  SKIP: {filename} — no extractable text")
                    continue

                documents.append(
                    Document(page_content=content, metadata={"source": filename})
                )
                logger.info(f"  PDF: {filename} — {len(reader.pages)} pages")

            # ── DOCX ────────────────────────────────────────
            elif ext == "docx":
                from docx import Document as DocxDoc

                doc = DocxDoc(filepath)
                content = "\n\n".join(
                    p.text for p in doc.paragraphs if p.text.strip()
                )

                if not content.strip():
                    logger.warning(f"  SKIP: {filename} — empty document")
                    continue

                documents.append(
                    Document(page_content=content, metadata={"source": filename})
                )
                logger.info(f"  DOCX: {filename} — loaded")

            # ── CSV ─────────────────────────────────────────
            elif ext == "csv":
                import csv

                with open(filepath, "r", encoding="utf-8") as f:
                    rows = list(csv.reader(f))
                content = "\n".join(", ".join(row) for row in rows)

                documents.append(
                    Document(page_content=content, metadata={"source": filename})
                )
                logger.info(f"  CSV: {filename} — {len(rows)} rows")

            # ── XLSX ────────────────────────────────────────
            elif ext == "xlsx":
                from openpyxl import load_workbook

                wb = load_workbook(filepath, read_only=True)
                parts = []
                for sheet in wb.sheetnames:
                    for row in wb[sheet].iter_rows(values_only=True):
                        vals = [str(c) if c is not None else "" for c in row]
                        parts.append(", ".join(vals))
                content = "\n".join(parts)

                documents.append(
                    Document(page_content=content, metadata={"source": filename})
                )
                logger.info(f"  XLSX: {filename} — {len(parts)} rows")

            else:
                logger.warning(f"  SKIP: {filename} — unsupported (.{ext})")

        except Exception as e:
            logger.error(f"  ERROR: {filename} — {e}")

    return documents


def main():
    config = load_config()
    data_dir = os.path.join(BASE_DIR, "data")
    vs_path = os.path.join(BASE_DIR, config.get("vector_store_path", "db/faiss_index"))

    mode = sys.argv[1].upper() if len(sys.argv) > 1 else "REPLACE"
    logger.info(f"Starting ingestion (mode={mode})...")

    # Load documents
    documents = load_documents(data_dir)
    if not documents:
        logger.warning("No documents found. Place files in data/ folder.")
        return

    logger.info(f"Loaded {len(documents)} document(s)")

    # Split into chunks
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.get("chunk_size", 1000),
        chunk_overlap=config.get("chunk_overlap", 200),
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")

    # Create embeddings & vector store
    from core_rag import create_embeddings

    embeddings = create_embeddings(config)

    from langchain_community.vectorstores import FAISS

    if mode == "APPEND" and os.path.exists(vs_path):
        logger.info("Appending to existing index...")
        vector_store = FAISS.load_local(
            vs_path, embeddings, allow_dangerous_deserialization=True
        )
        vector_store.add_documents(chunks)
    else:
        vector_store = FAISS.from_documents(chunks, embeddings)

    os.makedirs(os.path.dirname(vs_path), exist_ok=True)
    vector_store.save_local(vs_path)
    logger.info(f"Created! Total: {vector_store.index.ntotal} vectors")
    logger.info(f"Saved to: {vs_path}")


if __name__ == "__main__":
    main()
